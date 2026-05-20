"""
格式修改模块 - 使用 python-docx 直接修改 Word 文档格式
"""

import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def format_docx(input_path, options, output_path=None):
    """
    修改 Word 文档格式

    Args:
        input_path: Word 文档路径
        options: 格式选项字典
        output_path: 输出路径（可选）

    Returns:
        输出文件路径
    """
    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(f"输入文件不存在: {input_path}")

    # 确定输出路径
    if output_path is None:
        output_path = input_path.with_stem(input_path.stem + '_formatted')
    else:
        output_path = Path(output_path)

    # 备份原文件
    backup_path = input_path.with_suffix('.docx.backup')
    shutil.copy2(input_path, backup_path)

    # 打开文档
    doc = Document(input_path)

    # 应用格式
    _apply_format(doc, options)

    # 保存文档
    doc.save(output_path)

    print(f"✓ 格式修改成功: {input_path} → {output_path}")
    print(f"✓ 原文件已备份: {backup_path}")
    return output_path


def _apply_format(doc, options):
    """应用格式选项到文档"""

    # 修改正文字体和字号
    if 'font' in options or 'size' in options or 'color' in options:
        _modify_normal_style(doc, options)

    # 修改标题字体和字号
    if 'heading_font' in options or 'heading_size' in options or 'heading_color' in options:
        _modify_heading_styles(doc, options)

    # 修改行距
    if 'line_spacing' in options:
        _modify_line_spacing(doc, options['line_spacing'])

    # 修改页边距
    if any(k in options for k in ['margin_top', 'margin_bottom', 'margin_left', 'margin_right']):
        _modify_margins(doc, options)

    # 修改对齐方式
    if 'alignment' in options:
        _modify_alignment(doc, options['alignment'])

    # 修改首行缩进
    if 'first_line_indent' in options:
        _modify_first_line_indent(doc, options['first_line_indent'])


def _modify_normal_style(doc, options):
    """修改正文样式"""
    style = doc.styles['Normal']
    font = style.font

    if 'font' in options:
        font.name = options['font']
        # 设置中文字体
        from docx.oxml.ns import qn
        font.element.rPr.rFonts.set(qn('w:eastAsia'), options['font'])

    if 'size' in options:
        size = options['size']
        if isinstance(size, str):
            # 处理字号（如 "小四"、"五号"）
            size = _parse_chinese_size(size)
        font.size = Pt(size)

    if 'color' in options:
        color = options['color']
        if isinstance(color, str):
            color = _parse_color(color)
        font.color.rgb = color


def _modify_heading_styles(doc, options):
    """修改标题样式"""
    for i in range(1, 7):  # H1-H6
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            font = style.font

            if 'heading_font' in options:
                font.name = options['heading_font']
                from docx.oxml.ns import qn
                font.element.rPr.rFonts.set(qn('w:eastAsia'), options['heading_font'])

            if 'heading_size' in options:
                size = options['heading_size']
                if isinstance(size, str):
                    size = _parse_chinese_size(size)
                font.size = Pt(size)

            if 'heading_color' in options:
                color = options['heading_color']
                if isinstance(color, str):
                    color = _parse_color(color)
                font.color.rgb = color


def _modify_line_spacing(doc, line_spacing):
    """修改行距"""
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = line_spacing


def _modify_margins(doc, options):
    """修改页边距"""
    for section in doc.sections:
        if 'margin_top' in options:
            section.top_margin = _parse_length(options['margin_top'])
        if 'margin_bottom' in options:
            section.bottom_margin = _parse_length(options['margin_bottom'])
        if 'margin_left' in options:
            section.left_margin = _parse_length(options['margin_left'])
        if 'margin_right' in options:
            section.right_margin = _parse_length(options['margin_right'])


def _modify_alignment(doc, alignment):
    """修改对齐方式"""
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }

    if alignment in alignment_map:
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.alignment = alignment_map[alignment]


def _modify_first_line_indent(doc, indent):
    """修改首行缩进"""
    from docx.shared import Cm
    indent_value = _parse_length(indent)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.first_line_indent = indent_value


def _parse_length(value):
    """解析长度值（支持 cm、mm、in、pt）"""
    if isinstance(value, (int, float)):
        return Cm(value)

    value = str(value).strip()
    if value.endswith('cm'):
        return Cm(float(value[:-2]))
    elif value.endswith('mm'):
        return Cm(float(value[:-2]) / 10)
    elif value.endswith('in'):
        return Inches(float(value[:-2]))
    elif value.endswith('pt'):
        return Pt(float(value[:-2]))
    else:
        return Cm(float(value))


def _parse_chinese_size(size_str):
    """解析中文字号"""
    size_map = {
        '初号': 42,
        '小初': 36,
        '一号': 26,
        '小一': 24,
        '二号': 22,
        '小二': 18,
        '三号': 16,
        '小三': 15,
        '四号': 14,
        '小四': 12,
        '五号': 10.5,
        '小五': 9,
        '六号': 7.5,
        '小六': 6.5,
        '七号': 5.5,
        '八号': 5,
    }
    return size_map.get(size_str, 12)


def _parse_color(color_str):
    """解析颜色值"""
    color_map = {
        'red': RGBColor(0xFF, 0x00, 0x00),
        'blue': RGBColor(0x00, 0x00, 0xFF),
        'green': RGBColor(0x00, 0x80, 0x00),
        'black': RGBColor(0x00, 0x00, 0x00),
        'white': RGBColor(0xFF, 0xFF, 0xFF),
        'gray': RGBColor(0x80, 0x80, 0x80),
        'yellow': RGBColor(0xFF, 0xFF, 0x00),
        'orange': RGBColor(0xFF, 0xA5, 0x00),
        'purple': RGBColor(0x80, 0x00, 0x80),
    }

    if color_str.lower() in color_map:
        return color_map[color_str.lower()]

    # 尝试解析十六进制颜色值
    if color_str.startswith('#'):
        color_str = color_str[1:]
    if len(color_str) == 6:
        try:
            r = int(color_str[0:2], 16)
            g = int(color_str[2:4], 16)
            b = int(color_str[4:6], 16)
            return RGBColor(r, g, b)
        except ValueError:
            pass

    return RGBColor(0x00, 0x00, 0x00)  # 默认黑色


def apply_template(input_path, template_name, output_path=None):
    """
    应用模板格式

    Args:
        input_path: Word 文档路径
        template_name: 模板名称
        output_path: 输出路径（可选）

    Returns:
        输出文件路径
    """
    # 内置模板格式
    templates = {
        'thesis_cn': {
            'font': '宋体',
            'size': '小四',
            'heading_font': '黑体',
            'heading_size': '三号',
            'line_spacing': 1.5,
            'margin_top': '2.54cm',
            'margin_bottom': '2.54cm',
            'margin_left': '3.17cm',
            'margin_right': '3.17cm',
            'first_line_indent': '2cm',
        },
        'thesis_en': {
            'font': 'Times New Roman',
            'size': 12,
            'heading_font': 'Times New Roman',
            'heading_size': 16,
            'line_spacing': 2.0,
            'margin_top': '2.54cm',
            'margin_bottom': '2.54cm',
            'margin_left': '3.17cm',
            'margin_right': '3.17cm',
        },
        'resume': {
            'font': '微软雅黑',
            'size': '小四',
            'heading_font': '微软雅黑',
            'heading_size': '三号',
            'line_spacing': 1.25,
            'margin_top': '2cm',
            'margin_bottom': '2cm',
            'margin_left': '2.5cm',
            'margin_right': '2.5cm',
        },
    }

    if template_name not in templates:
        raise ValueError(f"未知模板: {template_name}，可用模板: {', '.join(templates.keys())}")

    return format_docx(input_path, templates[template_name], output_path)
