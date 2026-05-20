"""
测试文档转换模块
"""

import os
import tempfile
from pathlib import Path
import pytest

from src.converter import docx_to_markdown, markdown_to_docx


@pytest.fixture
def sample_docx():
    """创建测试用的 Word 文档"""
    from docx import Document

    doc = Document()
    doc.add_heading('测试文档', 0)
    doc.add_paragraph('这是一个测试段落。')
    doc.add_paragraph('这是第二个段落。')

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        yield Path(f.name)
        os.unlink(f.name)


@pytest.fixture
def sample_md():
    """创建测试用的 Markdown 文件"""
    content = """# 测试文档

这是一个测试段落。

这是第二个段落。
"""
    with tempfile.NamedTemporaryFile(suffix='.md', mode='w', delete=False) as f:
        f.write(content)
        yield Path(f.name)
        os.unlink(f.name)


def test_docx_to_markdown(sample_docx):
    """测试 Word 转 Markdown"""
    with tempfile.TemporaryDirectory() as tmpdir:
        output = Path(tmpdir) / 'test.md'
        result = docx_to_markdown(sample_docx, output)

        assert result.exists()
        assert result.suffix == '.md'
        content = result.read_text()
        assert '测试文档' in content


def test_markdown_to_docx(sample_md):
    """测试 Markdown 转 Word"""
    with tempfile.TemporaryDirectory() as tmpdir:
        output = Path(tmpdir) / 'test.docx'
        result = markdown_to_docx(sample_md, output_path=output)

        assert result.exists()
        assert result.suffix == '.docx'


def test_roundtrip(sample_docx):
    """测试往返转换（Word → Markdown → Word）"""
    with tempfile.TemporaryDirectory() as tmpdir:
        # Word → Markdown
        md_path = Path(tmpdir) / 'test.md'
        docx_to_markdown(sample_docx, md_path)

        # Markdown → Word
        docx_path = Path(tmpdir) / 'test_roundtrip.docx'
        markdown_to_docx(md_path, output_path=docx_path)

        assert docx_path.exists()
