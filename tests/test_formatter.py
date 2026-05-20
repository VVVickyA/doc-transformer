"""
测试格式修改模块
"""

import os
import tempfile
from pathlib import Path
import pytest

from src.formatter import format_docx, apply_template


@pytest.fixture
def sample_docx():
    """创建测试用的 Word 文档"""
    from docx import Document

    doc = Document()
    doc.add_heading('测试标题', 0)
    doc.add_paragraph('这是正文内容。')
    doc.add_paragraph('这是第二段。')

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        yield Path(f.name)
        os.unlink(f.name)


def test_format_docx(sample_docx):
    """测试格式修改"""
    with tempfile.TemporaryDirectory() as tmpdir:
        output = Path(tmpdir) / 'formatted.docx'
        options = {
            'font': '宋体',
            'size': 12,
            'line_spacing': 1.5,
        }
        result = format_docx(sample_docx, options, output)

        assert result.exists()
        assert result.suffix == '.docx'


def test_apply_template(sample_docx):
    """测试应用模板"""
    with tempfile.TemporaryDirectory() as tmpdir:
        output = Path(tmpdir) / 'thesis.docx'
        result = apply_template(sample_docx, 'thesis_cn', output)

        assert result.exists()
        assert result.suffix == '.docx'


def test_apply_template_invalid(sample_docx):
    """测试无效模板"""
    with pytest.raises(ValueError):
        apply_template(sample_docx, 'invalid_template')
